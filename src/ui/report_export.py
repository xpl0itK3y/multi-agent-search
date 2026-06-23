"""
report_export.py — PDF and DOCX export for research reports.

Usage:
    pdf_bytes  = generate_pdf(report_markdown, prompt, depth, created_at)
    docx_bytes = generate_docx(report_markdown, prompt, depth, created_at)
"""
from __future__ import annotations

import io
import os
import re
from urllib.parse import urlparse
from datetime import datetime
from typing import Optional

from src.ui.report_utils import clean_report as _clean_report_for_export  # noqa: F401


_LABELS: dict[str, dict[str, str]] = {
    "en": {
        "title":     "Research Report",
        "depth":     "Depth",
        "generated": "Generated",
        "page":      "Page",
    },
    "ru": {
        "title":     "Исследовательский отчёт",
        "depth":     "Глубина",
        "generated": "Создан",
        "page":      "Стр.",
    },
    "es": {
        "title":     "Informe de investigación",
        "depth":     "Profundidad",
        "generated": "Generado",
        "page":      "Pág.",
    },
}

# ──────────────────────────────────────────────────────────────────────────────
# Language detection
# ──────────────────────────────────────────────────────────────────────────────

def _detect_lang(text: str) -> str:
    """Detect prompt language: 'ru', 'es', or 'en'."""
    if not text:
        return "en"
    # Russian: significant Cyrillic ratio
    cyrillic = sum(1 for c in text if "Ѐ" <= c <= "ӿ")
    if cyrillic / len(text) > 0.15:
        return "ru"
    # Spanish: ñ / ¿ / ¡ are uniquely Spanish markers
    spanish = sum(1 for c in text if c in "ñÑ¿¡")
    if spanish >= 2 or (len(text) > 20 and spanish / len(text) > 0.003):
        return "es"
    return "en"


# ──────────────────────────────────────────────────────────────────────────────
# Font resolution (DejaVu → Unicode/Cyrillic; Helvetica → ASCII fallback)
# ──────────────────────────────────────────────────────────────────────────────

_DEJAVU_DIR = "/usr/share/fonts/truetype/dejavu"
_DEJAVU_REGULAR = "DejaVuSans.ttf"
_DEJAVU_BOLD    = "DejaVuSans-Bold.ttf"

_fonts_registered = False


def _register_pdf_fonts() -> dict[str, str]:
    """Register DejaVu TTF fonts if available; fall back to Helvetica."""
    global _fonts_registered
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.pdfmetrics import registerFontFamily
        from reportlab.pdfbase.ttfonts import TTFont

        r_path = os.path.join(_DEJAVU_DIR, _DEJAVU_REGULAR)
        b_path = os.path.join(_DEJAVU_DIR, _DEJAVU_BOLD)

        if os.path.isfile(r_path) and os.path.isfile(b_path):
            if not _fonts_registered:
                pdfmetrics.registerFont(TTFont("DV",   r_path))
                pdfmetrics.registerFont(TTFont("DV-B", b_path))
                # use regular as fallback for italic variants
                registerFontFamily("DV", normal="DV", bold="DV-B",
                                   italic="DV", boldItalic="DV-B")
                _fonts_registered = True
            return {"r": "DV", "b": "DV-B", "i": "DV", "bi": "DV-B"}
    except Exception:
        pass
    return {"r": "Helvetica", "b": "Helvetica-Bold",
            "i": "Helvetica-Oblique", "bi": "Helvetica-BoldOblique"}


# ──────────────────────────────────────────────────────────────────────────────
# Markdown parser
# ──────────────────────────────────────────────────────────────────────────────

def _link_display(label: str, url: str) -> str:
    """Return a short display string for a hyperlink.

    If the label is itself a URL (happens with legacy bare-URL source lines),
    show just the domain so the PDF column doesn't overflow.
    Otherwise cap at 70 chars.
    """
    if re.match(r"https?://", label):
        domain = urlparse(url).netloc.removeprefix("www.")
        return domain or label[:60]
    return (label[:70] + "…") if len(label) > 70 else label


_H2       = re.compile(r"^## (.+)$")
_H3       = re.compile(r"^### (.+)$")
_H4       = re.compile(r"^#### (.+)$")
_BULL     = re.compile(r"^[-*] (.+)$")
_ENUM     = re.compile(r"^\d+\. .+$")
_ENUM_CAP = re.compile(r"^\d+\. (.+)$")
_HR       = re.compile(r"^---+$")
_TABLE_ROW = re.compile(r"^\|.+\|$")
_TABLE_SEP = re.compile(r"^\|[-:| ]+\|$")


def _split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _parse_blocks(markdown: str) -> list[dict]:
    blocks: list[dict] = []
    lines = markdown.split("\n")
    i = 0
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        if m := _H2.match(stripped):
            blocks.append({"type": "h2", "text": m.group(1)})
        elif m := _H3.match(stripped):
            blocks.append({"type": "h3", "text": m.group(1)})
        elif m := _H4.match(stripped):
            blocks.append({"type": "h4", "text": m.group(1)})
        elif _HR.match(stripped):
            blocks.append({"type": "hr"})
        elif m := _BULL.match(stripped):
            items: list[str] = []
            while i < len(lines):
                bm = _BULL.match(lines[i].strip())
                if not bm:
                    break
                items.append(bm.group(1))
                i += 1
            blocks.append({"type": "bullets", "items": items})
            continue
        elif m := _ENUM_CAP.match(stripped):
            items = []
            while i < len(lines):
                em = _ENUM_CAP.match(lines[i].strip())
                if not em:
                    break
                items.append(em.group(1))
                i += 1
            blocks.append({"type": "ordered", "items": items})
            continue
        elif _TABLE_ROW.match(stripped):
            raw_rows: list[str] = []
            while i < len(lines):
                ln = lines[i].strip()
                if not _TABLE_ROW.match(ln):
                    break
                raw_rows.append(ln)
                i += 1
            # filter out separator rows (|---|---|)
            content_rows = [r for r in raw_rows if not _TABLE_SEP.match(r)]
            if len(content_rows) >= 1:
                headers = _split_table_row(content_rows[0])
                rows    = [_split_table_row(r) for r in content_rows[1:]]
                blocks.append({"type": "table", "headers": headers, "rows": rows})
            continue
        elif stripped:
            parts: list[str] = []
            while i < len(lines):
                ln = lines[i].strip()
                if (not ln or _H2.match(ln) or _H3.match(ln)
                        or _H4.match(ln) or _HR.match(ln)
                        or _BULL.match(ln) or _ENUM.match(ln)
                        or _TABLE_ROW.match(ln)):
                    break
                parts.append(lines[i])
                i += 1
            text = " ".join(p.strip() for p in parts if p.strip())
            if text:
                blocks.append({"type": "para", "text": text})
            continue
        i += 1
    return blocks


# ──────────────────────────────────────────────────────────────────────────────
# PDF  (reportlab)
# ──────────────────────────────────────────────────────────────────────────────

def generate_pdf(
    report: str,
    prompt: str,
    depth: str = "",
    created_at: Optional[str] = None,
) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        HRFlowable,
        ListFlowable,
        ListItem,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    report = _clean_report_for_export(report)
    lang  = _detect_lang(prompt)
    lbl   = _LABELS[lang]
    fonts = _register_pdf_fonts()
    F     = fonts["r"]
    FB    = fonts["b"]
    FI    = fonts["i"]
    FBI   = fonts["bi"]

    # ── palette ───────────────────────────────────────────────────────────────
    C_DARK    = colors.HexColor("#1e293b")
    C_BLUE    = colors.HexColor("#1d4ed8")
    C_BLUE_LT = colors.HexColor("#3b82f6")
    C_SLATE   = colors.HexColor("#334155")
    C_GRAY    = colors.HexColor("#64748b")
    C_LIGHT   = colors.HexColor("#f1f5f9")
    C_WHITE   = colors.white

    def S(name, **kw) -> ParagraphStyle:
        return ParagraphStyle(name, **kw)

    STYLES = {
        "h2":     S("h2",     fontName=FB,  fontSize=15,
                    textColor=C_BLUE,  leading=20, spaceAfter=4,  spaceBefore=14),
        "h3":     S("h3",     fontName=FB,  fontSize=12,
                    textColor=C_SLATE, leading=16, spaceAfter=3,  spaceBefore=10),
        "h4":     S("h4",     fontName=FBI, fontSize=10.5,
                    textColor=C_GRAY,  leading=14, spaceAfter=2,  spaceBefore=8),
        "para":   S("para",   fontName=F,   fontSize=10, leading=15,
                    textColor=C_DARK,  alignment=TA_JUSTIFY, spaceAfter=6),
        "bullet": S("bullet", fontName=F,   fontSize=10, leading=14,
                    textColor=C_DARK,  leftIndent=8, spaceAfter=2),
        "meta":   S("meta",   fontName=FI,  fontSize=9,
                    textColor=C_GRAY,  alignment=TA_LEFT),
        "title":  S("title",  fontName=FB,  fontSize=20,
                    textColor=C_WHITE, leading=26, alignment=TA_LEFT),
        "sub":    S("sub",    fontName=F,   fontSize=11,
                    textColor=C_WHITE, leading=15, alignment=TA_LEFT),
    }

    def _md_to_rl(text: str) -> str:
        text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        # unescape markdown bracket escapes written by _source_line()
        text = text.replace("\\[", "[").replace("\\]", "]")
        # markdown links [label](url) → clickable PDF link (display truncated)
        text = re.sub(
            r"\[([^\]]+)\]\((https?://[^\s)]+)\)",
            lambda m: f'<a href="{m.group(2)}" color="#1d4ed8">{_link_display(m.group(1), m.group(2))}</a>',
            text,
        )
        text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
        text = re.sub(
            r"\[S(\d+)\]",
            r'<font color="#1d4ed8" size="8"><super>[S\1]</super></font>',
            text,
        )
        return text

    # ── date string ───────────────────────────────────────────────────────────
    date_str = ""
    if created_at:
        try:
            dt = datetime.fromisoformat(created_at)
            date_str = dt.strftime("%d %b %Y")
        except Exception:
            date_str = created_at[:10]

    # ── header / footer ───────────────────────────────────────────────────────
    def _on_page(canvas, doc):
        canvas.saveState()
        W, H = A4
        canvas.setFillColor(C_BLUE)
        canvas.rect(0, H - 1.8*cm, W, 1.8*cm, fill=1, stroke=0)
        canvas.setFillColor(C_WHITE)
        canvas.setFont(FB, 9)
        canvas.drawString(2.5*cm, H - 1.15*cm, lbl["title"].upper())
        canvas.setFont(F, 8)
        prompt_short = (prompt[:80] + "…") if len(prompt) > 80 else prompt
        canvas.drawString(2.5*cm, H - 1.55*cm, prompt_short)
        if date_str:
            canvas.drawRightString(W - 2.5*cm, H - 1.35*cm, date_str)
        canvas.setFillColor(C_LIGHT)
        canvas.rect(0, 0, W, 1.1*cm, fill=1, stroke=0)
        canvas.setFillColor(C_GRAY)
        canvas.setFont(F, 8)
        canvas.drawCentredString(W / 2, 0.4*cm, f"{lbl['page']} {doc.page}")
        if depth:
            canvas.drawString(2.5*cm, 0.4*cm, f"{lbl['depth']}: {depth.upper()}")
        canvas.restoreState()

    # ── build story ───────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        topMargin=2.4*cm,
        bottomMargin=1.8*cm,
        leftMargin=2.5*cm,
        rightMargin=2.5*cm,
        onFirstPage=_on_page,
        onLaterPages=_on_page,
    )

    story = []

    # ── title block ───────────────────────────────────────────────────────────
    sub_text = prompt if len(prompt) <= 120 else prompt[:117] + "…"
    title_tbl = Table(
        [
            [Paragraph(lbl["title"], STYLES["title"])],
            [Paragraph(sub_text,     STYLES["sub"])],
        ],
        colWidths=[doc.width],
    )
    title_tbl.setStyle(TableStyle([
        ("BACKGROUND",   (0, 0), (-1, -1), C_BLUE),
        ("LEFTPADDING",  (0, 0), (-1, -1), 14),
        ("RIGHTPADDING", (0, 0), (-1, -1), 14),
        ("TOPPADDING",   (0, 0), (0, 0),   14),
        ("BOTTOMPADDING",(0,-1), (-1,-1),  14),
        ("ROWBACKGROUNDS",(0,0), (-1,-1),  [C_BLUE]),
    ]))
    story.append(title_tbl)

    # meta row
    meta_parts = []
    if depth:
        meta_parts.append(f"{lbl['depth']}: <b>{depth.upper()}</b>")
    if date_str:
        meta_parts.append(f"{lbl['generated']}: <b>{date_str}</b>")
    if meta_parts:
        story.append(Spacer(1, 6))
        story.append(Paragraph(" &nbsp;·&nbsp; ".join(meta_parts), STYLES["meta"]))

    story.append(Spacer(1, 10))
    story.append(HRFlowable(width="100%", thickness=1, color=C_LIGHT))
    story.append(Spacer(1, 8))

    # ── content ───────────────────────────────────────────────────────────────
    for block in _parse_blocks(report):
        btype = block["type"]

        if btype == "h2":
            if block["text"].lower() in ("sources", "источники"):
                story.append(Spacer(1, 6))
                story.append(HRFlowable(width="100%", thickness=0.5, color=C_LIGHT))
            story.append(Paragraph(_md_to_rl(block["text"]), STYLES["h2"]))

        elif btype == "h3":
            story.append(Paragraph(_md_to_rl(block["text"]), STYLES["h3"]))

        elif btype == "h4":
            story.append(Paragraph(_md_to_rl(block["text"]), STYLES["h4"]))

        elif btype == "hr":
            story.append(Spacer(1, 4))
            story.append(HRFlowable(width="100%", thickness=0.5, color=C_LIGHT))
            story.append(Spacer(1, 4))

        elif btype == "para":
            story.append(Paragraph(_md_to_rl(block["text"]), STYLES["para"]))

        elif btype == "bullets":
            items = [
                ListItem(
                    Paragraph(_md_to_rl(item), STYLES["bullet"]),
                    leftIndent=16,
                    bulletColor=C_BLUE_LT,
                )
                for item in block["items"]
            ]
            story.append(ListFlowable(items, bulletType="bullet",
                                      leftIndent=12, spaceBefore=2, spaceAfter=4))

        elif btype == "ordered":
            items = [
                ListItem(
                    Paragraph(_md_to_rl(item), STYLES["bullet"]),
                    leftIndent=16,
                    bulletColor=C_BLUE_LT,
                )
                for item in block["items"]
            ]
            story.append(ListFlowable(items, bulletType="1",
                                      leftIndent=12, spaceBefore=2, spaceAfter=4))

        elif btype == "table":
            headers  = block["headers"]
            tbl_rows = block["rows"]
            col_n    = max(len(headers), max((len(r) for r in tbl_rows), default=0), 1)

            def _pad(row: list[str], n: int) -> list[str]:
                return (row + [""] * n)[:n]

            S_TH = ParagraphStyle("th", fontName=FB, fontSize=9,
                                   textColor=C_WHITE, leading=12)
            S_TD = ParagraphStyle("td", fontName=F,  fontSize=9,
                                   textColor=C_DARK,  leading=12)

            tbl_data = [[Paragraph(_md_to_rl(h), S_TH) for h in _pad(headers, col_n)]]
            for row in tbl_rows:
                tbl_data.append([Paragraph(_md_to_rl(c), S_TD) for c in _pad(row, col_n)])

            col_w = doc.width / col_n
            pdf_tbl = Table(tbl_data, colWidths=[col_w] * col_n, repeatRows=1)
            pdf_tbl.setStyle(TableStyle([
                ("BACKGROUND",    (0, 0), (-1,  0),  C_BLUE),
                ("ROWBACKGROUNDS",(0, 1), (-1, -1),  [C_WHITE, C_LIGHT]),
                ("GRID",          (0, 0), (-1, -1),  0.4, C_GRAY),
                ("TOPPADDING",    (0, 0), (-1, -1),  4),
                ("BOTTOMPADDING", (0, 0), (-1, -1),  4),
                ("LEFTPADDING",   (0, 0), (-1, -1),  6),
                ("RIGHTPADDING",  (0, 0), (-1, -1),  6),
                ("VALIGN",        (0, 0), (-1, -1),  "TOP"),
            ]))
            story.append(Spacer(1, 6))
            story.append(pdf_tbl)
            story.append(Spacer(1, 8))

    doc.build(story)
    return buf.getvalue()


# ──────────────────────────────────────────────────────────────────────────────
# DOCX  (python-docx)
# ──────────────────────────────────────────────────────────────────────────────

def generate_docx(
    report: str,
    prompt: str,
    depth: str = "",
    created_at: Optional[str] = None,
) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    report = _clean_report_for_export(report)
    lang = _detect_lang(prompt)
    lbl  = _LABELS[lang]

    def hex_rgb(hex_str: str) -> RGBColor:
        h = hex_str.lstrip("#")
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    C_BLUE  = hex_rgb("1d4ed8")
    C_DARK  = hex_rgb("1e293b")
    C_SLATE = hex_rgb("334155")
    C_GRAY  = hex_rgb("64748b")

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── title block ───────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_para.add_run(lbl["title"])
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = C_BLUE
    title_para.paragraph_format.space_after = Pt(4)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub_run = sub.add_run(prompt if len(prompt) <= 140 else prompt[:137] + "…")
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = C_SLATE
    sub_run.font.italic = True
    sub.paragraph_format.space_after = Pt(4)

    # meta line
    date_str = ""
    if created_at:
        try:
            dt = datetime.fromisoformat(created_at)
            date_str = dt.strftime("%d %b %Y")
        except Exception:
            date_str = created_at[:10]

    meta_parts = []
    if depth:
        meta_parts.append(f"{lbl['depth']}: {depth.upper()}")
    if date_str:
        meta_parts.append(f"{lbl['generated']}: {date_str}")
    if meta_parts:
        meta_p = doc.add_paragraph(" · ".join(meta_parts))
        meta_p.runs[0].font.size = Pt(9)
        meta_p.runs[0].font.color.rgb = C_GRAY
        meta_p.paragraph_format.space_after = Pt(10)

    def _add_divider():
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "CBD5E1")
        pBdr.append(bottom)
        pPr.append(pBdr)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)

    _add_divider()

    # ── inline markdown helper ────────────────────────────────────────────────
    def _add_hyperlink(paragraph, url: str, label: str):
        """Add a clickable hyperlink run to paragraph."""
        r_id = paragraph.part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hl = OxmlElement("w:hyperlink")
        hl.set(qn("r:id"), r_id)
        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rStyle = OxmlElement("w:rStyle")
        rStyle.set(qn("w:val"), "Hyperlink")
        rPr.append(rStyle)
        new_run.append(rPr)
        t = OxmlElement("w:t")
        t.text = label
        new_run.append(t)
        hl.append(new_run)
        paragraph._p.append(hl)

    def _add_inline(paragraph, text: str):
        # unescape markdown bracket escapes from _source_line()
        text = text.replace("\\[", "[").replace("\\]", "]")
        # split on bold, citations, and markdown links
        _INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\[S\d+\]|\[[^\]]+\]\(https?://[^\s)]+\))")
        for part in _INLINE_RE.split(text):
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif re.match(r"\[S\d+\]", part):
                run = paragraph.add_run(part)
                run.font.color.rgb = C_BLUE
                run.font.size = Pt(8)
                rPr = run._r.get_or_add_rPr()
                vertAlign = OxmlElement("w:vertAlign")
                vertAlign.set(qn("w:val"), "superscript")
                rPr.append(vertAlign)
            elif m := re.match(r"\[([^\]]+)\]\((https?://[^\s)]+)\)", part):
                _add_hyperlink(paragraph, m.group(2), _link_display(m.group(1), m.group(2)))
            elif part:
                run = paragraph.add_run(part)
                run.font.color.rgb = C_DARK

    # ── content ───────────────────────────────────────────────────────────────
    for block in _parse_blocks(report):
        btype = block["type"]

        if btype == "h2":
            p = doc.add_heading("", level=1)
            p.clear()
            run = p.add_run(block["text"])
            run.font.bold = True
            run.font.size = Pt(15)
            run.font.color.rgb = C_BLUE
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)

        elif btype == "h3":
            p = doc.add_heading("", level=2)
            p.clear()
            run = p.add_run(block["text"])
            run.font.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = C_SLATE
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)

        elif btype == "h4":
            p = doc.add_heading("", level=3)
            p.clear()
            run = p.add_run(block["text"])
            run.font.bold = True
            run.font.italic = True
            run.font.size = Pt(11)
            run.font.color.rgb = C_GRAY
            p.paragraph_format.space_before = Pt(8)

        elif btype == "hr":
            _add_divider()

        elif btype == "para":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(6)
            _add_inline(p, block["text"])

        elif btype == "bullets":
            for item in block["items"]:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_after = Pt(2)
                _add_inline(p, item)

        elif btype == "ordered":
            for item in block["items"]:
                p = doc.add_paragraph(style="List Number")
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_after = Pt(2)
                _add_inline(p, item)

        elif btype == "table":
            headers  = block["headers"]
            tbl_rows = block["rows"]
            col_n    = max(len(headers), max((len(r) for r in tbl_rows), default=0), 1)

            def _pad(row: list[str], n: int) -> list[str]:
                return (row + [""] * n)[:n]

            tbl = doc.add_table(rows=1 + len(tbl_rows), cols=col_n)
            tbl.style = "Table Grid"

            # header row — blue background, white bold text
            for j, h in enumerate(_pad(headers, col_n)):
                cell = tbl.rows[0].cells[j]
                cell.text = ""
                run = cell.paragraphs[0].add_run(h)
                run.bold = True
                run.font.color.rgb = hex_rgb("FFFFFF")
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"),   "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"),  "1d4ed8")
                tc_pr.append(shd)

            # data rows — alternating fill
            for i_r, row in enumerate(tbl_rows):
                fill = "F1F5F9" if i_r % 2 == 0 else "FFFFFF"
                for j, cell_text in enumerate(_pad(row, col_n)):
                    cell = tbl.rows[i_r + 1].cells[j]
                    cell.text = ""
                    _add_inline(cell.paragraphs[0], cell_text)
                    tc_pr = cell._tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"),   "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"),  fill)
                    tc_pr.append(shd)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── HTML export (self-contained, shareable "mini-site") ──────────────────────────

import html as _html_mod

_MD_LINK = re.compile(r"\[([^\]]+)\]\((https?://[^)\s]+)\)")
_MD_BOLD = re.compile(r"\*\*(.+?)\*\*")
_MD_CITE = re.compile(r"\[S(\d+)\]")
_MD_HEADING = re.compile(r"^(#{1,4})\s+(.*)")
_MD_BULLET = re.compile(r"^[-*]\s+(.*)")
_MD_NUMBERED = re.compile(r"^\d+\.\s+(.*)")


def _md_inline(text: str) -> str:
    """Inline markdown on already HTML-escaped text: links, bold, [Sn] citations."""
    text = _MD_LINK.sub(r'<a href="\2" target="_blank" rel="noopener noreferrer">\1</a>', text)
    text = _MD_BOLD.sub(r"<strong>\1</strong>", text)
    text = _MD_CITE.sub(r'<sup class="cite">[S\1]</sup>', text)
    return text


def _slugify(text: str, used: set[str]) -> str:
    base = re.sub(r"[^\w\s-]", "", text, flags=re.UNICODE).strip().lower()
    base = re.sub(r"[\s_-]+", "-", base)[:60] or "section"
    slug, i = base, 2
    while slug in used:
        slug, i = f"{base}-{i}", i + 1
    used.add(slug)
    return slug


_MD_TABLE_SEP = re.compile(r"^\|?[\s:|-]+\|?$")


def _markdown_to_html(md: str) -> tuple[str, list[tuple[int, str, str]]]:
    """Markdown → HTML, also returning a (level, slug, text) outline for the table of contents.

    Headings get id anchors; the first paragraph is tagged as the lead; pipe tables render as
    real <table> (previously they leaked through as literal `| a | b |` paragraphs)."""
    out: list[str] = []
    para: list[str] = []
    list_type: Optional[str] = None
    toc: list[tuple[int, str, str]] = []
    used_slugs: set[str] = set()
    lead_done = False

    def flush_para() -> None:
        nonlocal lead_done
        if para:
            cls = "" if lead_done else ' class="lead"'
            out.append(f"<p{cls}>" + _md_inline(_html_mod.escape(" ".join(para))) + "</p>")
            lead_done = True
            para.clear()

    def close_list() -> None:
        nonlocal list_type
        if list_type:
            out.append(f"</{list_type}>")
            list_type = None

    lines = (md or "").split("\n")
    idx = 0
    while idx < len(lines):
        stripped = lines[idx].strip()
        nxt = lines[idx + 1].strip() if idx + 1 < len(lines) else ""
        if stripped.startswith("|") and "-" in nxt and _MD_TABLE_SEP.match(nxt):
            flush_para()
            close_list()
            header = _split_table_row(stripped)
            idx += 2  # consume header + separator rows
            rows: list[list[str]] = []
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                rows.append(_split_table_row(lines[idx].strip()))
                idx += 1
            thead = "".join(f"<th>{_md_inline(_html_mod.escape(c))}</th>" for c in header)
            tbody = "".join(
                "<tr>" + "".join(f"<td>{_md_inline(_html_mod.escape(c))}</td>" for c in r) + "</tr>"
                for r in rows
            )
            out.append(f'<div class="table-wrap"><table><thead><tr>{thead}</tr></thead><tbody>{tbody}</tbody></table></div>')
            continue
        if not stripped:
            flush_para()
            close_list()
            idx += 1
            continue
        heading = _MD_HEADING.match(stripped)
        if heading:
            flush_para()
            close_list()
            level = len(heading.group(1))
            text = heading.group(2)
            slug = _slugify(text, used_slugs)
            out.append(f'<h{level} id="{slug}">' + _md_inline(_html_mod.escape(text)) + f"</h{level}>")
            if level in (1, 2):
                toc.append((level, slug, text))
            idx += 1
            continue
        bullet = _MD_BULLET.match(stripped)
        numbered = _MD_NUMBERED.match(stripped)
        if bullet or numbered:
            flush_para()
            want = "ul" if bullet else "ol"
            if list_type != want:
                close_list()
                out.append(f"<{want}>")
                list_type = want
            out.append("<li>" + _md_inline(_html_mod.escape((bullet or numbered).group(1))) + "</li>")
            idx += 1
            continue
        close_list()
        para.append(stripped)
        idx += 1

    flush_para()
    close_list()
    return "\n".join(out), toc


_SANS = "-apple-system,BlinkMacSystemFont,'Segoe UI',Inter,system-ui,sans-serif"
_SERIF = "Georgia,'Iowan Old Style','Palatino Linotype','Times New Roman',serif"
_EDITORIAL = "'Iowan Old Style','Palatino Linotype',Palatino,Georgia,serif"
_STATUS_VARS = "--green:#10b981;--amber:#f59e0b;--red:#ef4444"

# Each theme is a block of CSS custom properties (the base rules below reference them), so a
# new look is just a new var set: bg/card/bd/ink/muted, an accent + accent2 (used for gradient
# flourishes), and a body/head font pairing. The user picks one in the export menu.
_THEME_VARS = {
    "light":     f"--bg:#faf9f5;--card:#ffffff;--bd:#eae7df;--ink:#2b2a27;--muted:#73706a;--accent:#d97757;--accent2:#c15f3c;--font-body:{_SANS};--font-head:{_SERIF}",
    "dark":      f"--bg:#1f1e1d;--card:#2a2927;--bd:#3a3835;--ink:#ece9e3;--muted:#a4a199;--accent:#e08a6a;--accent2:#d97757;--font-body:{_SANS};--font-head:{_SERIF}",
    "editorial": f"--bg:#f6f1e7;--card:#fffdf8;--bd:#e3dac6;--ink:#1f1b16;--muted:#6b6457;--accent:#9c4221;--accent2:#b45309;--font-body:{_EDITORIAL};--font-head:{_EDITORIAL}",
    "sepia":     f"--bg:#efe5d0;--card:#f7efde;--bd:#dac9a8;--ink:#3a2f21;--muted:#7c6c50;--accent:#a05a2c;--accent2:#7c4a2d;--font-body:{_EDITORIAL};--font-head:{_EDITORIAL}",
    "mono":      f"--bg:#ffffff;--card:#fafafa;--bd:#e3e3e3;--ink:#0f0f0f;--muted:#666666;--accent:#111111;--accent2:#555555;--font-body:{_EDITORIAL};--font-head:{_EDITORIAL}",
    "rose":      f"--bg:#fdf2f4;--card:#ffffff;--bd:#f4dde3;--ink:#3a2a2e;--muted:#8a6f76;--accent:#db2777;--accent2:#f472b6;--font-body:{_SANS};--font-head:{_SERIF}",
    "lavender":  f"--bg:#f5f3fb;--card:#ffffff;--bd:#e6e1f3;--ink:#2b2640;--muted:#79749a;--accent:#7c3aed;--accent2:#c026d3;--font-body:{_SANS};--font-head:{_SERIF}",
    "ocean":     f"--bg:#eef6f7;--card:#ffffff;--bd:#cee4e7;--ink:#14302f;--muted:#5b7a79;--accent:#0d9488;--accent2:#0ea5e9;--font-body:{_SANS};--font-head:{_SANS}",
    "slate":     f"--bg:#0f172a;--card:#1b2638;--bd:#2c3a52;--ink:#e8eef7;--muted:#93a3bb;--accent:#38bdf8;--accent2:#818cf8;--font-body:{_SANS};--font-head:{_SANS}",
    "midnight":  f"--bg:#13111c;--card:#1d1a2b;--bd:#2f2a42;--ink:#e9e6f5;--muted:#9b96b3;--accent:#a78bfa;--accent2:#f0abfc;--font-body:{_SANS};--font-head:{_SERIF}",
    "emerald":   f"--bg:#0b1411;--card:#13201b;--bd:#22332b;--ink:#e3efe8;--muted:#8fa89a;--accent:#34d399;--accent2:#a3e635;--font-body:{_SANS};--font-head:{_SERIF}",
    "forest":    f"--bg:#0d1410;--card:#16201a;--bd:#26342b;--ink:#e6efe5;--muted:#93a896;--accent:#84cc16;--accent2:#22c55e;--font-body:{_SANS};--font-head:{_SERIF}",
    "sunset":    f"--bg:#1a1014;--card:#251519;--bd:#3a2228;--ink:#f6e9e6;--muted:#bb9a93;--accent:#fb7185;--accent2:#fbbf24;--font-body:{_SANS};--font-head:{_SERIF}",
}

_HTML_BASE_CSS = """
*{box-sizing:border-box}
html{-webkit-text-size-adjust:100%;scroll-behavior:smooth}
body{margin:0;background:var(--bg);color:var(--ink);font:17px/1.75 var(--font-body);-webkit-font-smoothing:antialiased;text-rendering:optimizeLegibility}
::selection{background:color-mix(in srgb,var(--accent) 26%,transparent)}
.progress{position:fixed;top:0;left:0;height:3px;width:0;z-index:60;background:linear-gradient(90deg,var(--accent),var(--accent2))}
.totop{position:fixed;right:22px;bottom:22px;width:42px;height:42px;border:1px solid var(--bd);border-radius:50%;background:var(--card);color:var(--ink);font-size:18px;cursor:pointer;opacity:0;transform:translateY(8px);transition:opacity .2s,transform .2s,border-color .2s;z-index:60;box-shadow:0 6px 20px color-mix(in srgb,var(--ink) 12%,transparent)}
.totop.show{opacity:1;transform:none}
.totop:hover{border-color:var(--accent);color:var(--accent)}
.shell{max-width:1100px;margin:0 auto;padding:0 24px}
.layout{display:grid;gap:56px;grid-template-columns:minmax(0,1fr);align-items:start;padding:0 0 100px}
@media(min-width:1060px){.layout{grid-template-columns:minmax(0,1fr) 230px}}
main{min-width:0;width:100%;max-width:760px;padding-top:60px}
@media(min-width:1060px){main{justify-self:end}}
.eyebrow{display:inline-flex;align-items:center;gap:9px;font-size:12px;letter-spacing:.15em;text-transform:uppercase;color:var(--accent);font-weight:700}
.eyebrow::before{content:"";width:24px;height:2px;border-radius:2px;background:linear-gradient(90deg,var(--accent),var(--accent2))}
h1.title{font-family:var(--font-head);font-size:40px;line-height:1.12;letter-spacing:-.022em;margin:.34em 0 .26em}
.meta{color:var(--muted);font-size:13.5px;margin-bottom:34px}
.meta .dot{margin:0 9px;opacity:.5}
.cards{display:grid;grid-template-columns:repeat(auto-fit,minmax(132px,1fr));gap:12px;margin:0 0 16px}
.card{position:relative;overflow:hidden;background:var(--card);border:1px solid var(--bd);border-radius:16px;padding:16px;box-shadow:0 1px 2px color-mix(in srgb,var(--ink) 6%,transparent);transition:transform .15s ease,box-shadow .15s ease}
.card::before{content:"";position:absolute;top:0;bottom:0;left:0;width:3px;background:linear-gradient(var(--accent),var(--accent2));opacity:.9}
.card:hover{transform:translateY(-2px);box-shadow:0 10px 28px color-mix(in srgb,var(--ink) 12%,transparent)}
.card .lbl{font-size:11px;letter-spacing:.05em;text-transform:uppercase;color:var(--muted);font-weight:600}
.card .val{font-family:var(--font-head);font-size:30px;font-weight:700;line-height:1;margin-top:7px}
.card .sub{font-size:12px;color:var(--muted);margin-top:5px}
.g{color:var(--green)}.a{color:var(--amber)}.r{color:var(--red)}
.rail{order:2;padding-top:60px}
.toc{font-size:13.5px}
.toc summary{list-style:none;font-size:11px;letter-spacing:.12em;text-transform:uppercase;color:var(--muted);font-weight:700;margin-bottom:14px}
.toc summary::-webkit-details-marker{display:none}
.toc nav{display:flex;flex-direction:column}
.toc a{color:var(--muted);text-decoration:none;padding:5px 0 5px 14px;border-left:2px solid var(--bd);line-height:1.35;transition:color .15s,border-color .15s}
.toc a.lv2{padding-left:26px;font-size:12.5px}
.toc a:hover,.toc a.active{color:var(--accent);border-left-color:var(--accent)}
@media(min-width:1060px){.toc{position:sticky;top:30px}.toc summary{pointer-events:none}}
@media(max-width:1059px){
.rail{order:-1;padding-top:0;margin-bottom:6px}
.toc{background:var(--card);border:1px solid var(--bd);border-radius:14px;padding:14px 18px}
.toc summary{cursor:pointer;display:flex;align-items:center;justify-content:space-between;margin:0}
.toc summary::after{content:"⌄";font-size:17px;transition:transform .2s}
.toc[open] summary{margin-bottom:12px}.toc[open] summary::after{transform:rotate(180deg)}
}
article{font-size:17px}
article>:first-child{margin-top:0}
article h1,article h2,article h3,article h4{font-family:var(--font-head);line-height:1.25;letter-spacing:-.01em;margin:1.85em 0 .55em;color:var(--ink);scroll-margin-top:22px}
article h1{font-size:30px}
article h2{font-size:25px;margin-top:1.95em;padding-bottom:.3em;border-bottom:1px solid var(--bd)}
article h3{font-size:20px}
article h4{font-size:14px;color:var(--muted);text-transform:uppercase;letter-spacing:.04em}
article p{margin:.95em 0}
article p.lead{font-family:var(--font-head);font-size:21px;line-height:1.55;color:var(--ink);margin:.2em 0 1em}
article ul,article ol{margin:.8em 0;padding-left:1.4em}
article li{margin:.4em 0}article li::marker{color:var(--accent)}
article a{color:var(--accent);text-decoration:none;border-bottom:1px solid color-mix(in srgb,var(--accent) 35%,transparent);word-break:break-word}
article a:hover{border-bottom-color:var(--accent)}
article strong{font-weight:700;color:var(--ink)}
article blockquote{margin:1.4em 0;padding:.6em 1.3em;border-left:3px solid var(--accent);border-radius:0 12px 12px 0;background:color-mix(in srgb,var(--accent) 7%,transparent)}
article blockquote p{margin:.4em 0}
article code{font-family:ui-monospace,SFMono-Regular,Menlo,Consolas,monospace;font-size:.86em;background:color-mix(in srgb,var(--ink) 8%,transparent);padding:.12em .42em;border-radius:6px}
article pre{background:var(--card);border:1px solid var(--bd);border-radius:12px;padding:16px;overflow:auto;font-size:14px}
article pre code{background:none;padding:0}
.table-wrap{margin:1.5em 0;overflow-x:auto;border:1px solid var(--bd);border-radius:14px;box-shadow:0 1px 2px color-mix(in srgb,var(--ink) 5%,transparent)}
.table-wrap table{width:100%;border-collapse:collapse;font-size:14.5px;min-width:480px}
.table-wrap th,.table-wrap td{padding:11px 15px;text-align:left;border-bottom:1px solid var(--bd);vertical-align:top}
.table-wrap thead th{background:color-mix(in srgb,var(--accent) 11%,transparent);font-weight:700}
.table-wrap tbody tr:last-child td{border-bottom:0}
.table-wrap tbody tr:nth-child(even) td{background:color-mix(in srgb,var(--ink) 3%,transparent)}
sup.cite{display:inline-block;color:var(--accent);font-weight:700;font-size:.66em;vertical-align:super;padding:0 .12em;line-height:0}
hr{border:0;border-top:1px solid var(--bd);margin:2.6em 0}
footer{margin-top:60px;padding-top:24px;border-top:1px solid var(--bd);color:var(--muted);font-size:12.5px;display:flex;align-items:center;gap:9px}
footer::before{content:"";width:9px;height:9px;border-radius:50%;background:linear-gradient(135deg,var(--accent),var(--accent2))}
@media(max-width:600px){main{padding-top:40px}h1.title{font-size:30px}article h2{font-size:22px}article p.lead{font-size:19px}.shell{padding:0 18px}.layout{padding-bottom:64px}}
"""

_HEX_RE = re.compile(r"^#?[0-9a-fA-F]{3,8}$")


def _safe_color(value: Optional[str]) -> Optional[str]:
    """Accept only a hex color so a query param can't inject CSS."""
    if value and _HEX_RE.match(value.strip()):
        v = value.strip()
        return v if v.startswith("#") else f"#{v}"
    return None


def _build_css(theme: Optional[str], accent: Optional[str], base: Optional[str]) -> str:
    theme = (theme or "auto").lower()
    if theme == "custom":
        base_vars = _THEME_VARS["dark" if (base or "").lower() == "dark" else "light"]
        color = _safe_color(accent)
        if color:
            base_vars = re.sub(r"--accent:[^;]+", f"--accent:{color}", base_vars)
            base_vars = re.sub(r"--accent2:[^;]+", f"--accent2:{color}", base_vars)
        root = f":root{{{base_vars};{_STATUS_VARS}}}"
    elif theme == "auto":
        root = (
            f":root{{{_THEME_VARS['light']};{_STATUS_VARS}}}"
            f"@media (prefers-color-scheme:dark){{:root{{{_THEME_VARS['dark']}}}}}"
        )
    else:
        root = f":root{{{_THEME_VARS.get(theme, _THEME_VARS['light'])};{_STATUS_VARS}}}"
    return root + _HTML_BASE_CSS


def _card(label: str, value: str, sub: str = "", value_cls: str = "") -> str:
    cls = f" {value_cls}" if value_cls else ""
    sub_html = f'<div class="sub">{_html_mod.escape(sub)}</div>' if sub else ""
    return (
        f'<div class="card"><div class="lbl">{_html_mod.escape(label)}</div>'
        f'<div class="val{cls}">{_html_mod.escape(value)}</div>{sub_html}</div>'
    )


def _build_toc(toc: list[tuple[int, str, str]], labels: dict) -> str:
    entries = [(lvl, slug, text) for lvl, slug, text in toc if lvl in (1, 2)]
    if len(entries) < 3:
        return ""  # not worth a contents rail for a short report
    head = _html_mod.escape(labels.get("contents", "Contents"))
    links = "".join(
        f'<a class="lv{lvl}" href="#{slug}">{_html_mod.escape(text)}</a>'
        for lvl, slug, text in entries
    )
    return (
        f'<aside class="rail"><details class="toc" open>'
        f"<summary>{head}</summary><nav>{links}</nav></details></aside>"
    )


# Reading-progress bar, back-to-top toggle, and active-section highlighting in the TOC.
_TOC_SCRIPT = (
    "<script>(function(){"
    "var bar=document.querySelector('.progress'),top=document.querySelector('.totop');"
    "function s(){var h=document.documentElement,b=document.body,"
    "y=h.scrollTop||b.scrollTop,m=(h.scrollHeight||b.scrollHeight)-h.clientHeight;"
    "if(bar)bar.style.width=(m>0?y/m*100:0)+'%';if(top)top.classList.toggle('show',y>600);}"
    "document.addEventListener('scroll',s,{passive:true});s();"
    "if(top)top.addEventListener('click',function(){window.scrollTo({top:0,behavior:'smooth'});});"
    "var ls=[].slice.call(document.querySelectorAll('.toc a')),mp={};"
    "ls.forEach(function(a){mp[a.getAttribute('href').slice(1)]=a;});"
    "if('IntersectionObserver' in window){var o=new IntersectionObserver(function(es){"
    "es.forEach(function(e){if(e.isIntersecting){ls.forEach(function(a){a.classList.remove('active');});"
    "var a=mp[e.target.id];if(a)a.classList.add('active');}});},{rootMargin:'-12% 0px -78% 0px'});"
    "document.querySelectorAll('article h1[id],article h2[id]').forEach(function(h){o.observe(h);});}"
    "})();</script>"
)


def generate_html(
    report: str,
    prompt: str,
    depth: str = "",
    created_at: Optional[str] = None,
    scorecard: Optional[dict] = None,
    labels: Optional[dict] = None,
    theme: Optional[str] = None,
    accent: Optional[str] = None,
    base: Optional[str] = None,
) -> bytes:
    """Render the report as a self-contained, shareable HTML page (report + trust scorecard).

    ``theme`` picks a look (auto/light/dark/editorial/slate/custom); ``custom`` uses ``accent``
    (hex) over a light or dark ``base``.
    """
    labels = labels or {}
    when = created_at or ""
    try:
        when = datetime.fromisoformat(created_at).strftime("%Y-%m-%d") if created_at else ""
    except (ValueError, TypeError):
        pass

    cards = ""
    sc = scorecard or {}
    if sc.get("coverage_pct") is not None:
        cards += _card(labels.get("coverage", "Plan coverage"), f"{sc['coverage_pct']}%")
    if sc.get("integrity_pct") is not None:
        cls = "g" if sc["integrity_pct"] >= 80 else "a" if sc["integrity_pct"] >= 50 else "r"
        cards += _card(
            labels.get("citations", "Citations"), f"{sc['integrity_pct']}%",
            f"{sc.get('supported', 0)}/{sc.get('total', 0)}", cls,
        )
    if sc.get("sources"):
        high = sc.get("high_sources") or 0
        cards += _card(labels.get("sources", "Sources"), str(sc["sources"]),
                       f"{high} {labels.get('highQuality', 'high quality')}" if high else "")
    if sc.get("has_redteam"):
        cards += _card(
            labels.get("redteam", "Red-team"),
            f"{sc.get('challenged', 0)} / {sc.get('held', 0)}",
            labels.get("challengedHeld", "challenged / held"),
        )
    cards_html = f'<div class="cards">{cards}</div>' if cards else ""

    body, toc = _markdown_to_html(report)
    toc_html = _build_toc(toc, labels)
    words = len(re.findall(r"\w+", report, flags=re.UNICODE))
    reading = f"{max(1, round(words / 200))} {labels.get('readingTime', 'min read')}"
    meta_bits = [b for b in [depth, when, reading] if b]
    meta = '<span class="dot">·</span>'.join(_html_mod.escape(b) for b in meta_bits)
    footer = labels.get("footer", "Generated with verifiable research — every claim traceable to its source.")

    page = (
        "<!DOCTYPE html><html lang=\"ru\"><head><meta charset=\"utf-8\">"
        "<meta name=\"viewport\" content=\"width=device-width, initial-scale=1\">"
        f"<title>{_html_mod.escape(prompt[:120])}</title>"
        f"<style>{_build_css(theme, accent, base)}</style></head><body>"
        '<div class="progress"></div>'
        '<div class="shell"><div class="layout"><main>'
        f'<div class="eyebrow">{_html_mod.escape(labels.get("eyebrow", "Research report"))}</div>'
        f'<h1 class="title">{_html_mod.escape(prompt)}</h1>'
        f'<div class="meta">{meta}</div>'
        f"{cards_html}"
        f"<article>{body}</article>"
        f"<footer>{_html_mod.escape(footer)}</footer>"
        "</main>"
        f"{toc_html}"
        "</div></div>"
        '<button class="totop" aria-label="Top">↑</button>'
        f"{_TOC_SCRIPT}"
        "</body></html>"
    )
    return page.encode("utf-8")
